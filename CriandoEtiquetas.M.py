import os
import munip as mp
import docx as dx

print('-=-' * 20)
print('              Este programa Cria Etiquetas.')
print('-=-' * 20)

x = 'sim'
y = 'sim'
s = ('sim', 's')

valid_destino = False
valid_vara1 = False
lista_tipo = [1, 2, 3, 4]
lista_vara = ['1', '2', '3', '4', '5', '6', 'c']

while x in s:
    while valid_destino == False:
        print()
        tipo = input('Digite (1) para Vara Judicial, (2) para Vara Cível, (3) para Juizado Especial Cível, (4) para Distribuição da vara Cível: ')
        try:
            tipo = int(tipo)
            if tipo not in lista_tipo:
                print("===" * 20)
                print('Erro, Por favor use somente uma das opções informadas.')
                print("===" * 20)
            else:
                valid_destino = True
        except:
            print("===" * 20)
            print('Erro, Por favor use somente uma das opções informadas.')
            print("===" * 20)

    if tipo == 1:
        parte1 = "Vara Judicial da Comarca de"
        while valid_vara1 == False:
            print()
            vara = input("Especifique o numero da Vara ou digite 'c' caso não necessite de uma vara especifica: ")
            print()
            try:
                vara = str(vara)
                if vara not in lista_vara:
                    print("===" * 35)
                    print(
                        "Erro na entrada de dados, Ultilize somente numeros ou 'c' caso não precise de uma vara especifica.")
                    print("===" * 35)
                else:
                    valid_vara1 = True
            except:
                print("===" * 35)
                print(
                    "Erro na entrada de dados, Ultilize somente numeros ou 'c' caso não precise de uma vara especifica.")
                print("===" * 35)
        municipio = str(input('Digite a Comarca: ')).lower()
        enderecamento = mp.enderecos.get(municipio, 'Este município não costa no Banco de dados.')
        cepf = mp.cep.get(municipio, 'Este município não costa no Banco de dados')
        comarcaf = mp.comarca.get(municipio, 'Este município não costa no Banco de dados')
        if vara == 'c':
            primeiralinha = ('Aos Cuidados da {} {} - RS'.format(parte1, comarcaf))
        else:
            primeiralinha = ('Aos Cuidados da {}º {} {} - RS'.format(vara, parte1, comarcaf))

    elif tipo == 2:
        parte1 = "Vara Cível da Comarca de"
        while valid_vara1 == False:
            print()
            vara = input("Especifique o numero da Vara ou digite 'c' caso não necessite de uma vara especifica: ")
            print()
            try:
                vara = str(vara)
                if vara not in lista_vara:
                    print("===" * 35)
                    print(
                        "Erro na entrada de dados, Ultilize somente numeros ou 'c' caso não precise de uma vara especifica.")
                    print("===" * 35)
                else:
                    valid_vara1 = True
            except:
                print("===" * 35)
                print(
                    "Erro na entrada de dados, Ultilize somente numeros ou 'c' caso não precise de uma vara especifica.")
                print("===" * 35)
        municipio = str(input('Digite a Comarca: ')).lower()
        enderecamento = mp.enderecos.get(municipio, 'Este município não costa no Banco de dados.')
        cepf = mp.cep.get(municipio, 'Este município não costa no Banco de dados')
        comarcaf = mp.comarca.get(municipio, 'Este município não costa no Banco de dados')
        if vara == 'c':
            primeiralinha = ('Aos Cuidados da {} {} - RS'.format(parte1, comarcaf))
        else:
            primeiralinha = ('Aos Cuidados da {}º {} {} - RS'.format(vara, parte1, comarcaf))
    elif tipo == 3:
        parte1 = "Juizado especial Cível da Comarca de"
        municipio = str(input('Digite a Comarca: ')).lower()
        enderecamento = mp.enderecos.get(municipio, 'Este município não costa no Banco de dados.')
        cepf = mp.cep.get(municipio, 'Este município não costa no Banco de dados')
        comarcaf = mp.comarca.get(municipio, 'Este município não costa no Banco de dados')
        primeiralinha = ('Ao {} {} - RS'.format(parte1, comarcaf))
    elif tipo == 4:
        parte1 = "Distribuição junto à Vara Cível da Comarca de"
        municipio = str(input('Digite a Comarca: ')).lower()
        enderecamento = mp.enderecos.get(municipio, 'Este município não costa no Banco de dados.')
        cepf = mp.cep.get(municipio, 'Este município não costa no Banco de dados')
        comarcaf = mp.comarca.get(municipio, 'Este município não costa no Banco de dados')
        primeiralinha = ('{} {} - RS'.format(parte1, comarcaf))
    else:
        print('Erro, entre em contato com o Administrador.')

    print()
    print('-=-' * 30)
    print()
    print(primeiralinha)
    print('Endereço:', enderecamento)
    print('CEP:', cepf)
    print()
    print('-=-' * 30)
    print()

    y = input('Está correto?')
    print()

    if y in s:
        def escreverarquivo():
            document = dx.Document('test.docx')
            cepfinal = ('CEP: ', cepf)
            enderecamentofinal = ('Endereço: ', enderecamento)
            document.add_paragraph('===' * 20)
            document.add_paragraph(primeiralinha)
            document.add_paragraph(enderecamentofinal)
            document.add_paragraph(cepfinal)
            document.add_paragraph('===' * 20)
            document.add_paragraph('\n')
            document.save('test.docx')
    else:
        continue
    if tipo == 4:
        escreverarquivo()
        escreverarquivo()
    else:
        escreverarquivo()
    x = input('Deseja continuar? ').lower()

    valid_vara1 = False
    valid_destino = False


def abrirArquivo():
    os.startfile('test.docx')


abrirArquivo()
